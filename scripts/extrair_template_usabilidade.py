# -*- coding: utf-8 -*-
"""Extrai a estrutura do Widgets - Usabilidade.docx (template da doc de suporte)."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import RGBColor

doc = Document(r"D:\Trabalho\Twygo\Projetos\Widgets\Widgets - Usabilidade.docx")
print(f"# parágrafos: {len(doc.paragraphs)} | tabelas: {len(doc.tables)} | seções: {len(doc.sections)}")
print("\n--- ESTRUTURA (estilo | texto) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else "?"
    extra = ""
    if p.runs:
        r = p.runs[0]
        if r.bold: extra += " B"
        if r.font.size: extra += f" {r.font.size.pt}pt"
        if r.font.name: extra += f" {r.font.name}"
        if r.font.color and r.font.color.rgb: extra += f" #{r.font.color.rgb}"
    print(f"[{i:3d}] {style:18s}{extra:28s} | {t[:110]}")
for ti, tb in enumerate(doc.tables):
    print(f"\n--- TABELA {ti} ({len(tb.rows)}x{len(tb.columns)}) ---")
    for row in tb.rows[:6]:
        print("  | " + " | ".join(cell.text.strip()[:40] for cell in row.cells))
# imagens
import zipfile
with zipfile.ZipFile(r"D:\Trabalho\Twygo\Projetos\Widgets\Widgets - Usabilidade.docx") as z:
    imgs = [n for n in z.namelist() if n.startswith("word/media/")]
print(f"\n# imagens embutidas: {len(imgs)}")
