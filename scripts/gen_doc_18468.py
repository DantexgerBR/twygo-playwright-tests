"""Gera o documento de teste do 18468 como uma TABELA real (Word table), igual a
evidencia 'Quiz de Curiosidades Gamer': colunas #, Pergunta, Descricao, A, B, C, D,
Resposta Correta, Tipo, com ~100 linhas. Doc tabular faz a IA emitir uma TABELA
markdown no preview (o caminho de render onde o bug acontece), nao um outline.
Saida: evidencias/18468/quiz_grande.docx
"""
from pathlib import Path
from docx import Document

OUT = Path(__file__).resolve().parents[1] / "evidencias" / "18468"
OUT.mkdir(parents=True, exist_ok=True)

COLS = ["#", "Pergunta", "Descricao", "A", "B", "C", "D", "Resposta Correta", "Tipo"]

TEMAS = [
    ("Qual foi a inspiracao para o design do personagem Pac-Man?",
     ["Uma pizza sem uma fatia", "Uma bola de boliche", "Um volante de carro", "Um disco de vinil"], "A"),
    ("Em que ano surgiu o famoso Codigo Konami (cima, cima, baixo, baixo, tras, frente, tras, frente, B, A)?",
     ["1986", "1990", "1975", "2001"], "A"),
    ("Qual era a profissao original do Mario antes de ser encanador?",
     ["Carpinteiro", "Medico", "Cozinheiro", "Bombeiro"], "A"),
]

doc = Document()
doc.add_heading("Quiz de Curiosidades Gamer", level=1)
doc.add_paragraph("Importe esta tabela para gerar um questionario do tipo Prova.")

tab = doc.add_table(rows=1, cols=len(COLS))
tab.style = "Table Grid"
for i, h in enumerate(COLS):
    tab.rows[0].cells[i].text = h

for i in range(1, 101):
    perg, alts, correta = TEMAS[(i - 1) % len(TEMAS)]
    cells = tab.add_row().cells
    vals = [str(i), f"{perg} (item {i})", "", alts[0], alts[1], alts[2], alts[3], correta, "Unica escolha"]
    for j, v in enumerate(vals):
        cells[j].text = v

dest = OUT / "quiz_grande.docx"
doc.save(str(dest))
print(f"[ok] gerado {dest} ({dest.stat().st_size} bytes, tabela 100 linhas x {len(COLS)} colunas)")
