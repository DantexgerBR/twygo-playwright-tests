# -*- coding: utf-8 -*-
"""Gera 'Continuidade e sucessão - Usabilidade.docx' no padrão do template
Widgets - Usabilidade.docx (abre o template, limpa o corpo e reescreve —
herda Heading 1/2, fontes e cores).

v2: corrige bullet duplicado (add_paragraph inseria o parágrafo ANTES do
KeyError do estilo; agora o estilo é resolvido antes) e enxuga o texto pro
nível do template (doc de usabilidade, não especificação)."""
import sys

sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor

TEMPLATE = r"D:\Trabalho\Twygo\Projetos\Widgets\Widgets - Usabilidade.docx"
SAIDA = r"D:\Trabalho\Twygo\Projetos\Continuidade e sucessão\Continuidade e sucessão - Usabilidade.docx"
ROXO = RGBColor(0x78, 0x00, 0xFF)

doc = Document(TEMPLATE)

# limpa o corpo inteiro (mantém styles/sections)
body = doc.element.body
for el in list(body):
    if el.tag.endswith("}sectPr"):
        continue
    body.remove(el)

# resolve o estilo de bullet UMA vez (evita o bug do parágrafo duplicado)
try:
    BULLET_STYLE, BULLET_PREFIX = doc.styles["List Bullet"], ""
except KeyError:
    BULLET_STYLE, BULLET_PREFIX = doc.styles["List Paragraph"], "• "


def titulo(texto, tamanho=22):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(tamanho)
    r.font.name = "Baloo 2"
    r.font.color.rgb = ROXO
    return p


def par(texto=""):
    return doc.add_paragraph(texto)


def bullet(texto):
    return doc.add_paragraph(BULLET_PREFIX + texto, style=BULLET_STYLE)


def h1(texto):
    return doc.add_heading(texto, level=1)


def h2(texto):
    return doc.add_heading(texto, level=2)


def tabela(linhas, header=False):
    t = doc.add_table(rows=len(linhas), cols=len(linhas[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, linha in enumerate(linhas):
        for j, celula in enumerate(linha):
            t.cell(i, j).text = celula
            if header and i == 0:
                for r in t.cell(i, j).paragraphs[0].runs:
                    r.bold = True
    return t


# ============================== CONTEÚDO ==============================
titulo("Continuidade e sucessão — Documentação de Usabilidade")
par()
tabela([
    ["Objetivo: Descrever o funcionamento do módulo de Continuidade e sucessão e do widget de Tarefas, "
     "para apoio do time de suporte."],
    ["Localização na plataforma: Skills > Funções de negócio · menu Continuidade e sucessão (Dashboard "
     "geral, Análise individual, Ações de resposta, Parâmetros) · widget de Tarefas via Menu > Painéis."],
])
par()

titulo("Sumário", 16)
for s in ["O que é?", "Como funciona?", "Regras e comportamentos importantes", "Como configurar",
          "Fluxo de uso", "Telas do módulo", "Widget de Tarefas", "Visões por papel",
          "Comportamento importante"]:
    par(s)
par()

h1("O que é?")
par("Módulo que protege a organização da perda de pessoas-chave. O administrador cadastra funções "
    "críticas, informa quem as executa hoje e quem são os possíveis sucessores, e o sistema calcula um "
    "risco de descontinuidade de 0 a 100 para cada função.")
par("Para reduzir o risco, o usuário registra ações de resposta (ex.: capacitar um sucessor) e acompanha "
    "o efeito esperado em projeções de 6 e 12 meses, consolidadas no Dashboard geral.")

h1("Como funciona?")
par("Em Skills > Funções de negócio, cada função crítica é cadastrada em cinco abas: Identificação, "
    "Pessoas, Documentos, Competências e Continuidade (criticidade e mínimo de executores).")
par("O risco é calculado a partir de quatro fatores — probabilidade de perda dos executores, criticidade, "
    "força de trabalho (sucessores prontos) e cobertura — e recalculado automaticamente a cada alteração "
    "relevante. O link \"Como calculamos o risco?\" ao lado do score explica a fórmula.")
par("As ações de resposta não mudam o risco atual — mudam apenas as projeções de 6 e 12 meses. Quanto "
    "mais avançada a ação e mais próximo o prazo, maior o efeito projetado.")

h1("Regras e comportamentos importantes")
bullet("Funcionalidade liberada conforme contrato do ambiente — sem o módulo, a aba Continuidade fica "
       "desabilitada e as colunas de risco ficam ocultas.")
bullet("Mínimo de executores é obrigatório — sem ele o risco não é calculado.")
bullet("Criar uma ação não reduz o risco atual; o efeito aparece só nas projeções.")
bullet("Função sem executores assume probabilidade de perda máxima — o risco sobe.")
bullet("Ações canceladas saem das projeções; ações confidenciais só são vistas por Administrador e Líder.")
bullet("Toda alteração relevante gera histórico e snapshots diários (alimentam o gráfico do Dashboard).")

h1("Como configurar")
bullet("Skills > Funções de negócio: cadastrar as funções críticas (criticidade + mínimo de executores na "
       "aba Continuidade).")
bullet("Continuidade e sucessão > Parâmetros (só Administrador): cadastrar as iniciativas de resposta, com "
       "estratégia (Evitar, Mitigar, Transferir, Aceitar) e impactos no risco.")
bullet("Menu > Painéis: adicionar o widget de Tarefas a um painel e vincular ao Modo de uso — não há item "
       "\"Tarefas\" fixo no menu.")

h1("Fluxo de uso")
h2("1. Cadastrar a função crítica")
par("Acesse Skills > Funções de negócio e clique em + Adicionar. Informe nome, competências e, na aba "
    "Continuidade, a Criticidade e o Mínimo de executores.")
h2("2. Vincular executores e sucessores")
par("Na aba Pessoas, adicione executores com a Probabilidade de perda (Muito baixa a Alta) e sucessores "
    "com a Prontidão (Pronto agora, 6–12 meses, 12–24 meses, Futuro). A aba Sugestões indica as pessoas "
    "com maior aderência de competências.")
h2("3. Criar ações de resposta")
par("Em Ações de resposta, escolha função, estratégia, iniciativa, responsável, prazo e situação. O bloco "
    "\"Exposição ao risco\" mostra na hora a variação projetada.")
h2("4. Acompanhar")
par("Atualize a situação das ações (o sistema deriva o status real, ex.: \"Concluída com atraso\", "
    "\"Atrasada\") e acompanhe o consolidado no Dashboard geral. Cada responsável vê suas pendências no "
    "widget de Tarefas.")

h1("Telas do módulo")
tabela([
    ["Tela", "O que faz", "Perfil"],
    ["Dashboard geral", "Cards de risco e projeções, gráfico, tabelas de áreas/funções com maior risco, "
     "donuts de ações e mapa de risco de saída.", "Admin e Líder"],
    ["Análise individual", "Perfil consolidado por pessoa; probabilidade de perda editável na linha.",
     "Admin e Líder"],
    ["Ações de resposta", "Listagem e edição das ações de mitigação.", "Admin e Líder"],
    ["Parâmetros", "Iniciativas de resposta e seus impactos no risco.", "Somente Admin"],
    ["Skills > Funções de negócio", "Cadastro das funções críticas.", "Admin e Líder"],
    ["Widget de Tarefas", "Ações pendentes do usuário logado, por urgência.", "Líder e Aluno"],
], header=True)

h1("Widget de Tarefas")
par("Lista as ações pendentes em que o usuário logado é o responsável, agrupadas por urgência: "
    "Atrasadas → Hoje → Próximas. Concluídas e canceladas saem da lista automaticamente.")
bullet("Clicar em um card abre edição rápida: só Situação e Data de conclusão são editáveis; o link "
       "\"Abrir no módulo de origem\" leva à edição completa.")
bullet("Com mais tarefas do que cabe no widget, o rodapé mostra \"Ver todas (N)\" (abre a tela cheia).")
bullet("Sem pendências, exibe \"Nenhuma tarefa pendente.\".")
bullet("Aluno não vê ações confidenciais; Líder vê normalmente.")

h1("Visões por papel")
bullet("Administrador: acesso completo, sem filtros, e exclusividade na tela de Parâmetros.")
bullet("Líder de equipe: vê o módulo filtrado pelas pessoas que lidera; não vê Parâmetros (URL fora do "
       "escopo redireciona ao Dashboard).")
bullet("Aluno: interage apenas com o widget de Tarefas, quando seu Modo de uso dá acesso.")

h1("Comportamento importante")
bullet("Risco atual e projeções são coisas diferentes — ações só mexem nas projeções.")
bullet("Skills > Funções de negócio é a fonte única de funções; o módulo não duplica cadastro.")
bullet("Editar dados da pessoa em Análise individual sincroniza com o cadastro de usuários.")
bullet("Pontos em aberto nesta fase: comportamento do widget de Tarefas no tamanho mínimo e o botão "
       "\"Gerar com IA\" das competências (ainda não funcional).")

doc.save(SAIDA)
print(f"[ok] gerado: {SAIDA}")
