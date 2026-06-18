"""Gera um doc UNICO (nome+conteudo novos) pra forcar geracao FRESCA no agente
(que cacheia por documento). Tabela Word grande (100 linhas) com titulo e valores
unicos por timestamp. Imprime o caminho na ultima linha (stdout) pra o caller usar.
"""
import time
from pathlib import Path
from docx import Document

TS = time.strftime("%H%M%S")
OUT = Path(__file__).resolve().parents[1] / "evidencias" / "18468"
OUT.mkdir(parents=True, exist_ok=True)

COLS = ["#", "Pergunta", "A", "B", "C", "D", "Resposta Correta", "Tipo"]
doc = Document()
doc.add_heading(f"Quiz Unico {TS}", level=1)
tab = doc.add_table(rows=1, cols=len(COLS))
tab.style = "Table Grid"
for i, h in enumerate(COLS):
    tab.rows[0].cells[i].text = h

for i in range(1, 101):
    cells = tab.add_row().cells
    base = i * 7 + int(TS) % 1000   # valores unicos por execucao
    vals = [str(i), f"Qual o resultado da operacao numero {i} no lote {TS}?",
            str(base), str(base + 1), str(base + 2), str(base + 3), "A", "Unica escolha"]
    for j, v in enumerate(vals):
        cells[j].text = v

dest = OUT / f"quiz_unico_{TS}.docx"
doc.save(str(dest))
print(f"[ok] {dest.stat().st_size} bytes, tabela 100 linhas, titulo 'Quiz Unico {TS}'")
print(str(dest))
